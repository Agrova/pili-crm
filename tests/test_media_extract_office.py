"""Tests for analysis/media_extract/office.py — xlsx and docx parsers."""

from __future__ import annotations

from datetime import date
from pathlib import Path

import docx
import openpyxl
import pytest

from analysis.media_extract.office import (
    OfficeParseError,
    extract_docx,
    extract_xlsx,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def make_xlsx(path: Path, sheets: dict[str, list[list]]) -> Path:
    wb = openpyxl.Workbook()
    first = True
    for sheet_name, rows in sheets.items():
        if first:
            ws = wb.active
            ws.title = sheet_name
            first = False
        else:
            ws = wb.create_sheet(sheet_name)
        for row in rows:
            ws.append(row)
    wb.save(path)
    return path


def make_docx(path: Path, content: list[tuple[str, object]]) -> Path:
    document = docx.Document()
    for kind, data in content:
        if kind == "para":
            document.add_paragraph(str(data))
        elif kind == "table":
            rows_data: list[list[str]] = data  # type: ignore[assignment]
            table = document.add_table(rows=len(rows_data), cols=len(rows_data[0]))
            for i, row in enumerate(rows_data):
                for j, cell_text in enumerate(row):
                    table.cell(i, j).text = cell_text
    document.save(path)
    return path


# ---------------------------------------------------------------------------
# extract_xlsx
# ---------------------------------------------------------------------------


class TestExtractXlsx:
    def test_single_sheet(self, tmp_path: Path) -> None:
        path = make_xlsx(
            tmp_path / "orders.xlsx",
            {"Sheet1": [["Name", "Qty"], ["Oak beam", 5], ["Pine", 3]]},
        )
        result = extract_xlsx(path)
        lines = result.splitlines()
        assert lines[0] == "[Excel-файл: orders.xlsx]"
        assert 'Лист "Sheet1":' in result
        assert "Name\tQty" in result
        assert "Oak beam\t5" in result
        assert "Pine\t3" in result

    def test_multiple_sheets(self, tmp_path: Path) -> None:
        path = make_xlsx(
            tmp_path / "multi.xlsx",
            {
                "January": [["Item", "Price"], ["Saw", "1200"]],
                "February": [["Item", "Price"], ["Drill", "3400"]],
            },
        )
        result = extract_xlsx(path)
        assert 'Лист "January":' in result
        assert 'Лист "February":' in result
        assert result.index('Лист "January":') < result.index('Лист "February":')
        assert "Saw" in result
        assert "Drill" in result

    def test_empty_sheet(self, tmp_path: Path) -> None:
        path = make_xlsx(tmp_path / "empty.xlsx", {"EmptySheet": []})
        result = extract_xlsx(path)
        assert result.splitlines()[0] == "[Excel-файл: empty.xlsx]"
        assert 'Лист "EmptySheet": (пусто)' in result

    def test_mixed_types(self, tmp_path: Path) -> None:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Data"
        ws.append(["text_val", 42, None, date(2024, 1, 15)])
        wb.save(tmp_path / "mixed.xlsx")

        result = extract_xlsx(tmp_path / "mixed.xlsx")
        assert "text_val" in result
        assert "42" in result
        # None cell renders as empty string — row must still appear
        assert "Data" in result  # sheet name present
        # date str representation contains the year
        assert "2024" in result

    def test_file_not_found(self, tmp_path: Path) -> None:
        with pytest.raises(FileNotFoundError):
            extract_xlsx(tmp_path / "nonexistent.xlsx")

    def test_corrupted_file(self, tmp_path: Path) -> None:
        bad = tmp_path / "bad.xlsx"
        bad.write_bytes(b"this is definitely not a valid xlsx zip archive")
        with pytest.raises(OfficeParseError):
            extract_xlsx(bad)

    # optional
    def test_unicode(self, tmp_path: Path) -> None:
        path = make_xlsx(
            tmp_path / "unicode.xlsx",
            {"Лист": [["Привет", "мир", "🪵"]]},
        )
        result = extract_xlsx(path)
        assert "Привет" in result
        assert "мир" in result
        assert "🪵" in result


# ---------------------------------------------------------------------------
# extract_docx
# ---------------------------------------------------------------------------


class TestExtractDocx:
    def test_simple_paragraphs(self, tmp_path: Path) -> None:
        path = make_docx(
            tmp_path / "simple.docx",
            [
                ("para", "First paragraph"),
                ("para", "Second paragraph"),
                ("para", "Third paragraph"),
            ],
        )
        result = extract_docx(path)
        lines = result.splitlines()
        assert lines[0] == "[Word-файл: simple.docx]"
        assert "First paragraph" in result
        assert "Second paragraph" in result
        assert "Third paragraph" in result
        # order preserved
        assert result.index("First") < result.index("Second") < result.index("Third")

    def test_with_table(self, tmp_path: Path) -> None:
        path = make_docx(
            tmp_path / "table.docx",
            [
                ("para", "Before table"),
                ("table", [["Cell A", "Cell B"], ["Cell C", "Cell D"]]),
                ("para", "After table"),
            ],
        )
        result = extract_docx(path)
        assert "Before table" in result
        assert "Cell A\tCell B" in result
        assert "Cell C\tCell D" in result
        assert "After table" in result
        assert result.index("Before table") < result.index("Cell A") < result.index("After table")

    def test_empty_paragraphs_skipped(self, tmp_path: Path) -> None:
        document = docx.Document()
        document.add_paragraph("")
        document.add_paragraph("Visible text")
        document.add_paragraph("   ")  # whitespace-only
        document.add_paragraph("")
        path = tmp_path / "empties.docx"
        document.save(path)

        result = extract_docx(path)
        assert "Visible text" in result
        # no blank lines in output (header line + content only)
        output_lines = result.splitlines()
        assert all(line != "" for line in output_lines)

    def test_file_not_found(self, tmp_path: Path) -> None:
        with pytest.raises(FileNotFoundError):
            extract_docx(tmp_path / "nonexistent.docx")

    def test_corrupted_file(self, tmp_path: Path) -> None:
        bad = tmp_path / "bad.docx"
        bad.write_bytes(b"this is not a valid docx zip archive at all")
        with pytest.raises(OfficeParseError):
            extract_docx(bad)

    # optional
    def test_unicode(self, tmp_path: Path) -> None:
        path = make_docx(
            tmp_path / "unicode.docx",
            [("para", "Кириллица и эмоджи 🪵🔨")],
        )
        result = extract_docx(path)
        assert "Кириллица" in result
        assert "🪵" in result
